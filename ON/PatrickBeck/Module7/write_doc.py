import docx
import requests


res = requests.get('http://www.textfiles.com/computers/mrdos1.txt') #This site was down on multiple attempts when trying to access it
res.raise_for_status()
# Why 406? just because to add first few paragraphs.
# Can be any number depending on how much you want to print.
text = res.text[:9310]
doc = docx.Document()

# add 'This is the first paragraph.' to doc and store it to para1
para1 = doc.add_paragraph('This is the first paragraph.', 'Title')
# add text from the link above to the document
for p in text.split('\n\n'):
    # print('-', p)
    try:
        doc.add_paragraph(p)
    except ValueError:
        continue

doc.save('firstWord.docx')