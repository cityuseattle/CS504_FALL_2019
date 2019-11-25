import docx 
import requests

res = requests.get('http://www.textfiles.com/computers/mrdos1.txt')
res.raise_for_status()

text = res.text[:9310]
doc = docx.Document()

para1 = doc.add_paragraph('This is the first paragraph.', 'Title')
for p in text.split('\n\n'):
    try:
        doc.add_paragraph(p)
    except ValueError:
        continue
    
doc.save('firstWord.docx')